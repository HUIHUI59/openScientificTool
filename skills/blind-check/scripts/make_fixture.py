#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_fixture.py — 生成一份「埋了 6 类身份泄露」的测试稿(fixture.docx + fixture.pdf),
用来验证/演示 blind_check.py 能不能全部检出。人物与项目号均为虚构。

依赖: pip install python-docx pypdf
产物: 本目录下 fixture.docx(含①属性②修订③批注⑤致谢基金⑥自引) + fixture.pdf(含④元数据)
"""
import os
import shutil
import sys
import zipfile

from docx import Document
from pypdf import PdfWriter

for _s in ("stdout", "stderr"):
    try:
        getattr(sys, _s).reconfigure(encoding="utf-8")
    except Exception:
        pass

BASE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(BASE, "fixture.docx")
PDF = os.path.join(BASE, "fixture.pdf")

# --- docx: ①文件属性 ⑤致谢/基金 ⑥自引 ---
doc = Document()
doc.core_properties.author = "Zhang Wei"
doc.core_properties.last_modified_by = "张伟"
p1 = doc.add_paragraph("本研究基于问卷调查数据展开分析。")
doc.add_paragraph("笔者此前的研究（张伟，2023）表明社会资本与主观幸福感正相关。")
doc.add_paragraph("致谢：感谢导师李强教授的悉心指导。本研究受国家社会科学基金项目（21BSH123）资助。")

# --- ③批注署名(python-docx >= 1.2 支持) ---
try:
    doc.add_comment(runs=p1.runs, text="这里建议补充样本量说明", author="张伟", initials="ZW")
    print("批注已写入(python-docx)")
except Exception as e:
    print("批注写入失败(python-docx 版本过低?):", e)
doc.save(DOCX)

# --- ②修订记录署名: 直接往 document.xml 里注入一条 w:ins ---
tmp = DOCX + ".tmp"
with zipfile.ZipFile(DOCX, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == "word/document.xml":
            xml = data.decode("utf-8")
            ins = ('<w:ins w:id="99" w:author="Zhang Wei" w:date="2026-07-01T00:00:00Z">'
                   '<w:r><w:t>（修订插入的句子）</w:t></w:r></w:ins>')
            xml = xml.replace("</w:body>", "<w:p>" + ins + "</w:p></w:body>", 1)
            data = xml.encode("utf-8")
        zout.writestr(item, data)
shutil.move(tmp, DOCX)
print("fixture.docx 已生成:", DOCX)

# --- pdf: ④PDF 元数据 ---
w = PdfWriter()
w.add_blank_page(width=595, height=842)
w.add_metadata({"/Author": "Zhang Wei", "/Creator": "Microsoft Word - 张伟学位论文v3.docx"})
with open(PDF, "wb") as f:
    w.write(f)
print("fixture.pdf 已生成:", PDF)
